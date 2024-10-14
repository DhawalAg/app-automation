import os
from docx import Document
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Database IDs
COMPANY_TRACKER_ID = os.environ["COMPANY_TRACKER_ID"]
APPLICATION_TRACKER_ID = os.environ["APPLICATION_TRACKER_ID"]

BASE_DIR = os.environ["ALL_APPLICATIONS_DIRECTORY"]

def create_folder_if_not_exists(folder_path: str):
    """
    Create a folder if it doesn't exist.
    """
    if not os.path.exists(folder_path):
        os.makedirs(folder_path)
        print(f"Folder created: {folder_path}")
    else:
        print(f"Folder already exists: {folder_path}")

def create_word_document(job_description: str, file_path: str):
    """
    Create a Word document with the provided job description and save it to the specified file path.
    """
    doc = Document()
    doc.add_heading("Job Description", level=1)
    doc.add_paragraph(job_description)
    doc.save(file_path)
    print(f"Word document saved: {file_path}")

def create_folders_and_file(role_name: str, company_name: str, job_description: str, base_path: str = BASE_DIR):
    """
    Create the necessary folder structure and save the job description in a Word document.
    
    Parameters:
    - role_name: Name of the role (folder will be created for this role).
    - company_name: Name of the company (top-level folder).
    - job_description: Job description to be saved in the Word document.
    - base_path: Base directory to start from (default is the current directory).
    """
    # Define the path for the company folder
    company_folder_path = os.path.join(base_path, company_name)
    create_folder_if_not_exists(company_folder_path)
    
    # Define the path for the role folder within the company folder
    role_folder_path = os.path.join(company_folder_path, role_name)
    create_folder_if_not_exists(role_folder_path)
    
    # Define the path for the Word document to be saved inside the role folder
    word_file_path = os.path.join(role_folder_path, f"Job_Description.docx")
    
    # Create the Word document with the job description
    create_word_document(job_description, word_file_path)

def main():
    # Input from the user
    role_name = input("Enter the role name: ")
    company_name = input("Enter the company name: ")
    job_description = input("Enter the job description: ")
    
    # Call the function to create folders and save the job description
    create_folders_and_file(role_name, company_name, job_description)

if __name__ == "__main__":
    main()